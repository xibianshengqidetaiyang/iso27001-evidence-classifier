
#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ISO/IEC 27001 证据智能分类脚本（多对多）
------------------------------------------------
特点：
1. 同时看文件名 + 文本内容，不依赖证据文件名完全规范
2. 一个证据可命中多个条款/控制项（many-to-many）
3. 自动从你的 ISO 27001 审核手册中提取规则：控制编号、标题、必填证据、核心关键词、辅助关键词
4. 输出命中原因、命中关键词、命中片段、分数、置信度
5. 支持 md/txt/docx/pdf/xlsx/csv/json 等常见证据格式

适用场景：
- ISO27001 / TISAX 资料初筛
- 审核前资料归档
- 规则库雏形
- AI/脚本辅助审核

用法示例：
python iso27001_evidence_classifier.py \
  --evidence-dir ./evidence \
  --output-dir ./output \
  --handbooks ./ISO27001_2022_信息安全管理体系总则审核手册_V1.0.md ./ISO27001_2022_附录A_A5-A8_控制项审核手册_V1.0.md

若不传 --handbooks，会优先尝试使用脚本所在目录下的两个默认手册文件名。
"""
from __future__ import annotations

import argparse
import csv
import dataclasses
import json
import math
import os
import re
import sys
import textwrap
from collections import Counter, defaultdict
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Sequence, Set, Tuple

try:
    import docx  # python-docx
except Exception:
    docx = None

try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None

try:
    import openpyxl
except Exception:
    openpyxl = None

# -----------------------------
# 可调参数
# -----------------------------
MAX_TEXT_CHARS = 200_000
SNIPPET_WINDOW = 36
DEFAULT_THRESHOLD = 3.5
TOP_K_PER_DOC = 12

GENERIC_TERMS = {
    # 很泛的中文词
    "适用范围", "责任人", "批准人", "发布日期", "版本号", "记录", "审批", "复核记录", "台账", "日志", "工单",
    "项目", "系统范围", "安全要求", "更新日期", "触发条件", "流程", "程序", "制度", "标准", "控制",
    "作者", "审核", "审查", "日期", "编号", "范围", "职责", "owner", "record", "approval", "version",
    "review", "procedure", "policy", "standard", "evidence", "log", "table", "sheet", "form",
    "开发", "测试", "设计", "能力", "沟通", "独立", "责任", "保持", "人员", "工具", "检索", "使用", "应用", "程序", "总则", "管理", "资源", "流程", "计划", "方针", "事件", "权限", "风险",
    # 手册中经常重复出现的审核字段
    "整改闭环", "例外审批", "会议纪要", "台账日志工单", "记录编号", "处置时点", "项目系统范围", "测试评审记录",
    "交叉验证", "现场观察", "正式受控版本", "适用性状态", "实施说明", "排除理由",
}

# 这些词不是完全泛化，但很像“证据类型词”，用于辅助评分而非强绑定某个控制
EVIDENCE_TYPE_TERMS = {
    "soa", "声明适用性", "适用性声明",
    "制度", "程序", "标准", "方针", "办法", "规范", "流程", "清单", "台账", "报告", "记录", "日志", "工单",
    "计划", "清册", "矩阵", "清单", "纪要", "台账", "审批", "截图", "配置", "基线", "策略", "协议", "合同",
    "预案", "演练", "培训", "签到", "考核", "资产", "风险", "权限", "账号", "漏洞", "备份", "变更", "测试",
    "拓扑", "网络", "供应商", "个人信息", "隐私", "事件", "威胁情报", "审计", "评审",
}

STOPWORDS_EN = {
    "the", "and", "or", "for", "of", "to", "in", "on", "with", "from", "a", "an", "is", "are",
    "within", "during", "use", "using", "management", "security", "information",
}

# -----------------------------
# 数据结构
# -----------------------------
@dataclasses.dataclass
class Rule:
    control_id: str
    title_raw: str
    title_cn: str
    title_en: str
    required_evidence: List[str]
    core_keywords: List[str]
    aux_keywords: List[str]
    domain_terms: List[str]
    evidence_type_terms: List[str]


@dataclasses.dataclass
class DocumentInfo:
    path: Path
    file_name: str
    suffix: str
    text: str
    normalized_text: str
    compact_text: str
    compact_name: str


@dataclasses.dataclass
class MatchResult:
    file_path: str
    control_id: str
    title_cn: str
    title_en: str
    score: float
    confidence: str
    matched_in_name: List[str]
    matched_domain_terms: List[str]
    matched_core_keywords: List[str]
    matched_aux_keywords: List[str]
    matched_required_evidence: List[str]
    matched_evidence_type_terms: List[str]
    snippets: List[str]
    reasons: List[str]


# -----------------------------
# 文本工具
# -----------------------------
def normalize_text(text: str) -> str:
    text = text.lower()
    text = text.replace("\u3000", " ")
    text = re.sub(r"[“”\"'`*]+", " ", text)
    text = re.sub(r"[（）()【】\[\]<>《》]", " ", text)
    text = re.sub(r"[_/\\|,:;，。；：、\-+*=~!@#$%^&?]+", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def compact_text(text: str) -> str:
    text = normalize_text(text)
    text = re.sub(r"\s+", "", text)
    return text


def unique_keep_order(items: Iterable[str]) -> List[str]:
    seen: Set[str] = set()
    out: List[str] = []
    for item in items:
        v = item.strip()
        if not v:
            continue
        if v not in seen:
            seen.add(v)
            out.append(v)
    return out


def split_aliases(text: str) -> List[str]:
    """把一条证据/关键词拆成多种可匹配别名"""
    text = text.strip()
    if not text:
        return []
    aliases = [text]

    # 常见斜杠、括号展开
    for part in re.split(r"[；;，,、]|(?:\s{2,})", text):
        part = part.strip(" -")
        if part and part != text:
            aliases.append(part)

    # 中英括号
    paren_parts = re.findall(r"[（(]([^()（）]{1,80})[)）]", text)
    aliases.extend(paren_parts)

    # 斜杠形式，如 责任人/Owner
    slash_parts = re.split(r"[/｜|]", text)
    aliases.extend([p.strip() for p in slash_parts if p.strip()])

    # 去掉“最近12个月”“正式受控版本”等修饰，得到更核心短语
    simplified = text
    simplified = re.sub(r"最近\d+个月", "", simplified)
    simplified = re.sub(r"正式受控版本", "", simplified)
    simplified = re.sub(r"的", "", simplified)
    simplified = re.sub(r"相关", "", simplified)
    simplified = re.sub(r"中该控制", "", simplified)
    simplified = re.sub(r"适用性状态|实施说明|排除理由", "", simplified)
    simplified = re.sub(r"\s+", "", simplified)
    if simplified and simplified != text:
        aliases.append(simplified)

    # 再次粗拆
    rough_parts = re.split(r"[ /（）()【】\[\]-]", simplified)
    aliases.extend([p.strip() for p in rough_parts if len(p.strip()) >= 2])

    # 英文词拆分
    norm = normalize_text(text)
    en_words = [w for w in re.split(r"\s+", norm) if len(w) >= 4 and w not in STOPWORDS_EN]
    aliases.extend(en_words)

    # 过滤过泛短词
    cleaned = []
    for a in aliases:
        a = a.strip().lower()
        a = a.replace("（", "").replace("）", "")
        if not a:
            continue
        if len(a) == 1:
            continue
        if a in {"owner", "record", "approval"}:
            cleaned.append(a)
            continue
        if a in GENERIC_TERMS and len(a) <= 4:
            continue
        cleaned.append(a)
    return unique_keep_order(cleaned)


def contains_phrase(compact_haystack: str, phrase: str) -> bool:
    p = compact_text(phrase)
    return bool(p) and p in compact_haystack


def contains_control_id(normalized_haystack: str, control_id: str) -> bool:
    pattern = re.compile(rf"(?<![\w.]){re.escape(control_id.lower())}(?![\w.])")
    return bool(pattern.search(normalized_haystack))


def first_snippet(normalized_text_input: str, phrase: str, window: int = SNIPPET_WINDOW) -> Optional[str]:
    """从已规范化文本中取片段，避免重复重算。"""
    if not normalized_text_input or not phrase:
        return None
    target = normalize_text(phrase)
    if not target:
        return None
    idx = normalized_text_input.find(target)
    if idx == -1:
        return phrase
    start = max(0, idx - window)
    end = min(len(normalized_text_input), idx + len(target) + window)
    snippet = normalized_text_input[start:end].strip()
    return snippet


def confidence_from_score(score: float) -> str:
    if score >= 9:
        return "高"
    if score >= 6:
        return "中"
    return "低"


# -----------------------------
# 规则解析
# -----------------------------
CONTROL_RE = re.compile(r"^###\s+((?:A\.\d+\.\d+)|(?:[4-9]\.\d(?:\.\d+)?)|(?:10\.\d(?:\.\d+)?))\s+(.+?)\s*$")

def split_title(title_raw: str) -> Tuple[str, str]:
    title_raw = title_raw.strip()
    # 例如：Threat intelligence（威胁情报）
    m1 = re.match(r"(.+?)（(.+?)）", title_raw)
    if m1:
        left, right = m1.group(1).strip(), m1.group(2).strip()
        # 猜哪个是中文
        if re.search(r"[\u4e00-\u9fff]", left):
            return left, right
        return right, left
    # 例如：理解组织及其环境（Understanding the organization and its context）
    m2 = re.match(r"(.+?)\((.+?)\)", title_raw)
    if m2:
        left, right = m2.group(1).strip(), m2.group(2).strip()
        if re.search(r"[\u4e00-\u9fff]", left):
            return left, right
        return right, left
    # fallback
    if re.search(r"[\u4e00-\u9fff]", title_raw):
        return title_raw, ""
    return "", title_raw


def parse_keyword_line(line: str, marker: str) -> List[str]:
    # 例如：- **核心关键词（缺失直接不合规）**：A.5.7；威胁情报；适用范围；...
    if marker not in line:
        return []
    _, right = line.split("：", 1) if "：" in line else line.split(":", 1)
    parts = [re.sub(r"[*`]+", "", p).strip(" -") for p in re.split(r"[；;]", right) if p.strip(" -")]
    return parts


def parse_required_evidence(block_lines: List[str], start_idx: int) -> List[str]:
    out = []
    for i in range(start_idx + 1, len(block_lines)):
        line = block_lines[i].strip()
        if not line:
            continue
        if "【刚性校验关键词/字段清单】" in line:
            break
        if line.startswith("-"):
            out.append(re.sub(r"[*`]+", "", line.lstrip("- ").strip()))
    return out


def derive_domain_terms(control_id: str, title_cn: str, title_en: str,
                        required_evidence: List[str], core_keywords: List[str], aux_keywords: List[str]) -> List[str]:
    terms = []

    # 标题及标题拆分
    for src in [title_cn, title_en]:
        if not src:
            continue
        terms.append(src)
        for x in split_aliases(src):
            terms.append(x)

    # 从核心关键词和必填证据中提取更“专属”的短语
    for src in (core_keywords + required_evidence):
        for x in split_aliases(src):
            if x in GENERIC_TERMS:
                continue
            if len(x) >= 2:
                terms.append(x)

    # 一些辅助关键词有明显领域特征
    for src in aux_keywords:
        for x in split_aliases(src):
            if x in GENERIC_TERMS:
                continue
            if len(x) >= 3:
                terms.append(x)

    # 去掉太泛词
    filtered = []
    for t in unique_keep_order(terms):
        nt = normalize_text(t)
        if not nt:
            continue
        if nt in GENERIC_TERMS:
            continue
        if len(compact_text(nt)) <= 1:
            continue
        filtered.append(t)
    return filtered


def derive_evidence_type_terms(required_evidence: List[str], core_keywords: List[str], aux_keywords: List[str]) -> List[str]:
    text = " ".join(required_evidence + core_keywords + aux_keywords)
    terms = []
    for t in EVIDENCE_TYPE_TERMS:
        if compact_text(t) and compact_text(t) in compact_text(text):
            terms.append(t)
    return unique_keep_order(terms)


def parse_handbook(md_path: Path) -> List[Rule]:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    controls: List[Tuple[str, str, List[str]]] = []

    current_id = None
    current_title = None
    current_block: List[str] = []
    for line in lines:
        m = CONTROL_RE.match(line.strip())
        if m:
            if current_id is not None:
                controls.append((current_id, current_title, current_block))
            current_id = m.group(1).strip()
            current_title = m.group(2).strip()
            current_block = [line]
        else:
            if current_id is not None:
                current_block.append(line)
    if current_id is not None:
        controls.append((current_id, current_title, current_block))

    rules: List[Rule] = []
    for cid, raw_title, block in controls:
        title_cn, title_en = split_title(raw_title)
        required = []
        core = []
        aux = []

        for i, line in enumerate(block):
            s = line.strip()
            if "【一票否决必填证据清单】" in s:
                required.extend(parse_required_evidence(block, i))
            if "核心关键词（" in s and ("：" in s or ":" in s):
                core.extend(parse_keyword_line(s, "核心关键词"))
            if "辅助关键词（" in s and ("：" in s or ":" in s):
                aux.extend(parse_keyword_line(s, "辅助关键词"))

        required = unique_keep_order(required)
        core = unique_keep_order(core)
        aux = unique_keep_order(aux)
        domain_terms = derive_domain_terms(cid, title_cn, title_en, required, core, aux)
        evidence_type_terms = derive_evidence_type_terms(required, core, aux)

        rules.append(
            Rule(
                control_id=cid,
                title_raw=raw_title,
                title_cn=title_cn,
                title_en=title_en,
                required_evidence=required,
                core_keywords=core,
                aux_keywords=aux,
                domain_terms=domain_terms,
                evidence_type_terms=evidence_type_terms,
            )
        )
    return rules


def load_rules_from_handbooks(handbook_paths: Sequence[Path]) -> List[Rule]:
    rules = []
    for p in handbook_paths:
        if not p.exists():
            raise FileNotFoundError(f"找不到手册文件：{p}")
        rules.extend(parse_handbook(p))
    # 去重：同 control_id 以更丰富内容为准
    best: Dict[str, Rule] = {}
    for r in rules:
        prev = best.get(r.control_id)
        if prev is None:
            best[r.control_id] = r
            continue
        prev_score = len(prev.required_evidence) + len(prev.core_keywords) + len(prev.aux_keywords)
        cur_score = len(r.required_evidence) + len(r.core_keywords) + len(r.aux_keywords)
        if cur_score > prev_score:
            best[r.control_id] = r
    return [best[k] for k in sorted(best.keys(), key=sort_control_id)]


def sort_control_id(cid: str):
    if cid.startswith("A."):
        a, b = cid.split(".")[1:]
        return (100, int(a), int(b))
    parts = cid.split(".")
    nums = tuple(int(x) for x in parts)
    return (0,) + nums


# -----------------------------
# 文档读取
# -----------------------------
def read_text_file(path: Path) -> str:
    for enc in ("utf-8", "utf-8-sig", "gb18030", "latin-1"):
        try:
            return path.read_text(encoding=enc)
        except Exception:
            continue
    return ""


def read_docx_file(path: Path) -> str:
    if docx is None:
        return ""
    try:
        d = docx.Document(str(path))
        paras = [p.text for p in d.paragraphs if p.text.strip()]
        # 表格也很重要
        for table in d.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paras.append(row_text)
        return "\n".join(paras)
    except Exception:
        return ""


def read_pdf_file(path: Path) -> str:
    if PdfReader is None:
        return ""
    try:
        reader = PdfReader(str(path))
        texts = []
        for page in reader.pages[:80]:
            try:
                texts.append(page.extract_text() or "")
            except Exception:
                continue
        return "\n".join(texts)
    except Exception:
        return ""


def read_xlsx_file(path: Path) -> str:
    if openpyxl is None:
        return ""
    try:
        wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
        texts = []
        for ws in wb.worksheets[:15]:
            texts.append(f"[sheet]{ws.title}")
            row_count = 0
            for row in ws.iter_rows(values_only=True):
                vals = [str(v).strip() for v in row if v is not None and str(v).strip()]
                if vals:
                    texts.append(" | ".join(vals))
                    row_count += 1
                if row_count >= 1500:
                    break
        return "\n".join(texts)
    except Exception:
        return ""


def read_csv_file(path: Path) -> str:
    try:
        texts = []
        with path.open("r", encoding="utf-8-sig", newline="") as f:
            reader = csv.reader(f)
            for i, row in enumerate(reader):
                vals = [str(v).strip() for v in row if str(v).strip()]
                if vals:
                    texts.append(" | ".join(vals))
                if i >= 3000:
                    break
        return "\n".join(texts)
    except Exception:
        return read_text_file(path)


def read_json_file(path: Path) -> str:
    try:
        obj = json.loads(path.read_text(encoding="utf-8"))
        return json.dumps(obj, ensure_ascii=False, indent=2)
    except Exception:
        return read_text_file(path)


def read_document(path: Path) -> DocumentInfo:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".log", ".ini", ".conf", ".yaml", ".yml", ".xml", ".html"}:
        text = read_text_file(path)
    elif suffix == ".docx":
        text = read_docx_file(path)
    elif suffix == ".pdf":
        text = read_pdf_file(path)
    elif suffix in {".xlsx", ".xlsm"}:
        text = read_xlsx_file(path)
    elif suffix == ".csv":
        text = read_csv_file(path)
    elif suffix == ".json":
        text = read_json_file(path)
    else:
        text = read_text_file(path)

    text = text[:MAX_TEXT_CHARS]
    normalized = normalize_text(text)
    return DocumentInfo(
        path=path,
        file_name=path.name,
        suffix=suffix,
        text=text,
        normalized_text=normalized,
        compact_text=compact_text(text),
        compact_name=compact_text(path.name),
    )


def iter_evidence_files(evidence_dir: Path) -> Iterable[Path]:
    for path in evidence_dir.rglob("*"):
        if path.is_file() and not path.name.startswith("~$"):
            yield path


# -----------------------------
# 分类打分
# -----------------------------
def compute_keyword_df(rules: Sequence[Rule]) -> Dict[str, int]:
    df = Counter()
    for r in rules:
        seen = set()
        for item in r.domain_terms + r.core_keywords + r.aux_keywords + r.required_evidence:
            for alias in split_aliases(item):
                key = compact_text(alias)
                if key and key not in seen and len(key) >= 2:
                    seen.add(key)
        for key in seen:
            df[key] += 1
    return dict(df)


def idf_weight(alias: str, df_map: Dict[str, int], n_rules: int, floor: float = 0.35) -> float:
    key = compact_text(alias)
    if not key:
        return floor
    df = df_map.get(key, n_rules)
    val = math.log((n_rules + 1) / (df + 1)) + 0.6
    return max(floor, round(val, 3))


def match_aliases(target_compact: str, source_items: Sequence[str]) -> List[str]:
    hits = []
    for item in source_items:
        for alias in split_aliases(item):
            if contains_phrase(target_compact, alias):
                hits.append(alias)
                break
    return unique_keep_order(hits)


def infer_evidence_terms(doc: DocumentInfo) -> List[str]:
    pool = doc.file_name + " " + doc.text[:15000]
    pool_compact = compact_text(pool)
    hits = []
    for term in EVIDENCE_TYPE_TERMS:
        if contains_phrase(pool_compact, term):
            hits.append(term)
    return unique_keep_order(hits)



def strong_signal_hits(hits: Sequence[str], df_map: Dict[str, int], n_rules: int) -> List[str]:
    out = []
    for h in hits:
        c = compact_text(h)
        if not c:
            continue
        w = idf_weight(h, df_map, n_rules)
        # 既考虑词长，也考虑稀有度；短但非常稀有的词也可作为强信号
        if len(c) >= 4 and w >= 0.75:
            out.append(h)
        elif len(c) >= 3 and w >= 1.55:
            out.append(h)
    return unique_keep_order(out)


def classify_document(doc: DocumentInfo, rules: Sequence[Rule], df_map: Dict[str, int], threshold: float) -> List[MatchResult]:
    results = []
    n_rules = max(1, len(rules))
    doc_evidence_terms = infer_evidence_terms(doc)

    for rule in rules:
        score = 0.0
        reasons = []
        snippets = []

        matched_in_name = match_aliases(doc.compact_name, [rule.title_cn, rule.title_en] + rule.domain_terms)
        matched_domain = match_aliases(doc.compact_text, rule.domain_terms)
        matched_core = match_aliases(doc.compact_text, rule.core_keywords)
        matched_aux = match_aliases(doc.compact_text, rule.aux_keywords)
        matched_req = match_aliases(doc.compact_text, rule.required_evidence)
        matched_type = [t for t in doc_evidence_terms if t in rule.evidence_type_terms]

        # 1) 文件名命中：很值钱
        if matched_in_name:
            add = min(4.2, 1.2 * len(matched_in_name))
            score += add
            reasons.append(f"文件名命中 {len(matched_in_name)} 个控制相关词")
            for hit in matched_in_name[:3]:
                snippets.append(f"[文件名] {hit}")

        # 2) 领域词命中：最重要
        domain_score = 0.0
        for hit in matched_domain:
            w = idf_weight(hit, df_map, n_rules)
            domain_score += 1.6 * w
        if matched_domain:
            domain_score = min(domain_score, 8.0)
            score += domain_score
            reasons.append(f"正文命中领域词 {len(matched_domain)} 个")

        # 3) 核心关键词命中：高权重，但做 IDF 降泛化
        core_score = 0.0
        for hit in matched_core:
            w = idf_weight(hit, df_map, n_rules)
            core_score += 1.2 * w
        if matched_core:
            core_score = min(core_score, 6.0)
            score += core_score
            reasons.append(f"正文命中核心关键词 {len(matched_core)} 个")

        # 4) 必填证据描述命中：说明文档内容更靠近审核证据
        req_score = 0.0
        for hit in matched_req:
            w = idf_weight(hit, df_map, n_rules)
            req_score += 1.0 * w
        if matched_req:
            req_score = min(req_score, 5.0)
            score += req_score
            reasons.append(f"正文命中必填证据描述 {len(matched_req)} 个")

        # 5) 辅助关键词：补充分
        aux_score = 0.0
        for hit in matched_aux:
            w = idf_weight(hit, df_map, n_rules, floor=0.25)
            aux_score += 0.6 * w
        if matched_aux:
            aux_score = min(aux_score, 3.0)
            score += aux_score
            reasons.append(f"正文命中辅助关键词 {len(matched_aux)} 个")

        # 6) 文档像某类证据
        if matched_type:
            add = min(2.0, 0.55 * len(matched_type))
            score += add
            reasons.append(f"文档特征像 {len(matched_type)} 类证据")
            for t in matched_type[:3]:
                snippets.append(f"[证据类型] {t}")

        # 7) 如果出现控制编号，直接加分
        if contains_control_id(doc.normalized_text, rule.control_id):
            score += 2.4
            reasons.append("正文直接出现控制编号")

        # 8) 如果标题中文/英文完整出现，额外加分
        for title in [rule.title_cn, rule.title_en]:
            if title and contains_phrase(doc.compact_text, title):
                score += 1.8
                reasons.append("正文直接出现控制标题")

        title_hit = bool(
            (rule.title_cn and len(compact_text(rule.title_cn)) >= 4 and contains_phrase(doc.compact_text, rule.title_cn))
            or
            (rule.title_en and len(compact_text(rule.title_en)) >= 6 and contains_phrase(doc.compact_text, rule.title_en))
        )
        direct_id_hit = contains_control_id(doc.normalized_text, rule.control_id)
        strong_hits = strong_signal_hits(matched_domain + matched_req + matched_core, df_map, n_rules)

        if score >= threshold and (matched_in_name or title_hit or direct_id_hit or strong_hits):
            sample_terms = strong_hits[:3] + matched_req[:2] + matched_core[:2] + matched_in_name[:1]
            for term in sample_terms:
                sp = first_snippet(doc.normalized_text, term)
                if sp:
                    snippets.append(sp)

            result = MatchResult(
                file_path=str(doc.path),
                control_id=rule.control_id,
                title_cn=rule.title_cn,
                title_en=rule.title_en,
                score=round(score, 3),
                confidence=confidence_from_score(score),
                matched_in_name=matched_in_name,
                matched_domain_terms=matched_domain,
                matched_core_keywords=matched_core,
                matched_aux_keywords=matched_aux,
                matched_required_evidence=matched_req,
                matched_evidence_type_terms=matched_type,
                snippets=unique_keep_order(snippets)[:8],
                reasons=reasons,
            )
            results.append(result)

    results.sort(key=lambda x: (-x.score, x.control_id))
    return results[:TOP_K_PER_DOC]


# -----------------------------
# 输出
# -----------------------------
def write_results(output_dir: Path, rules: Sequence[Rule], docs: Sequence[DocumentInfo], all_results: Dict[str, List[MatchResult]]) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)

    # 1) 明细 CSV
    detail_csv = output_dir / "classification_detail.csv"
    with detail_csv.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerow([
            "file_path", "control_id", "title_cn", "title_en", "score", "confidence",
            "matched_in_name", "matched_domain_terms", "matched_core_keywords", "matched_aux_keywords",
            "matched_required_evidence", "matched_evidence_type_terms", "reasons", "snippets"
        ])
        for file_key, matches in all_results.items():
            for m in matches:
                writer.writerow([
                    m.file_path, m.control_id, m.title_cn, m.title_en, m.score, m.confidence,
                    " | ".join(m.matched_in_name),
                    " | ".join(m.matched_domain_terms),
                    " | ".join(m.matched_core_keywords),
                    " | ".join(m.matched_aux_keywords),
                    " | ".join(m.matched_required_evidence),
                    " | ".join(m.matched_evidence_type_terms),
                    "；".join(m.reasons),
                    " || ".join(m.snippets),
                ])

    # 2) 按文档汇总 JSON
    summary_json = output_dir / "classification_summary.json"
    summary = []
    for doc in docs:
        matches = all_results.get(str(doc.path), [])
        summary.append({
            "file_path": str(doc.path),
            "top_controls": [
                {
                    "control_id": m.control_id,
                    "title_cn": m.title_cn,
                    "title_en": m.title_en,
                    "score": m.score,
                    "confidence": m.confidence,
                    "reasons": m.reasons,
                    "matched_domain_terms": m.matched_domain_terms,
                    "matched_core_keywords": m.matched_core_keywords,
                    "matched_required_evidence": m.matched_required_evidence,
                    "snippets": m.snippets,
                }
                for m in matches
            ]
        })
    summary_json.write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")

    # 3) 按控制项聚合 CSV
    by_control: Dict[str, List[MatchResult]] = defaultdict(list)
    for matches in all_results.values():
        for m in matches:
            by_control[m.control_id].append(m)

    control_csv = output_dir / "classification_by_control.csv"
    with control_csv.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["control_id", "title_cn", "title_en", "file_count", "files"])
        rule_map = {r.control_id: r for r in rules}
        for cid in sorted(by_control.keys(), key=sort_control_id):
            rule = rule_map[cid]
            files = sorted({m.file_path for m in by_control[cid]})
            writer.writerow([cid, rule.title_cn, rule.title_en, len(files), " | ".join(files)])

    # 4) 编译后的规则 CSV，方便你继续调规则
    rules_csv = output_dir / "compiled_rules.csv"
    with rules_csv.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerow([
            "control_id", "title_cn", "title_en",
            "required_evidence", "core_keywords", "aux_keywords", "domain_terms", "evidence_type_terms"
        ])
        for r in rules:
            writer.writerow([
                r.control_id, r.title_cn, r.title_en,
                " | ".join(r.required_evidence),
                " | ".join(r.core_keywords),
                " | ".join(r.aux_keywords),
                " | ".join(r.domain_terms),
                " | ".join(r.evidence_type_terms),
            ])

    # 5) 简短报告
    report_md = output_dir / "report.md"
    lines = []
    lines.append("# ISO27001 证据分类报告")
    lines.append("")
    lines.append(f"- 处理文件数：{len(docs)}")
    lines.append(f"- 命中分类文件数：{sum(1 for v in all_results.values() if v)}")
    lines.append(f"- 规则数：{len(rules)}")
    lines.append("")
    lines.append("## 每个文件的 Top 结果")
    lines.append("")
    for doc in docs:
        lines.append(f"### {doc.path.name}")
        matches = all_results.get(str(doc.path), [])
        if not matches:
            lines.append("- 未达到阈值，建议人工复核或降低阈值。")
            lines.append("")
            continue
        for m in matches[:6]:
            lines.append(f"- **{m.control_id} {m.title_cn or m.title_en}** | 分数={m.score} | 置信度={m.confidence}")
            if m.matched_domain_terms:
                lines.append(f"  - 领域词：{', '.join(m.matched_domain_terms[:8])}")
            if m.matched_required_evidence:
                lines.append(f"  - 必填证据描述：{', '.join(m.matched_required_evidence[:6])}")
            if m.snippets:
                lines.append(f"  - 片段：{m.snippets[0]}")
        lines.append("")
    report_md.write_text("\n".join(lines), encoding="utf-8")


# -----------------------------
# CLI
# -----------------------------
def default_handbook_paths(script_dir: Path) -> List[Path]:
    return [
        script_dir / "ISO27001_2022_信息安全管理体系总则审核手册_V1.0.md",
        script_dir / "ISO27001_2022_附录A_A5-A8_控制项审核手册_V1.0.md",
    ]


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="ISO27001 证据多标签分类脚本")
    parser.add_argument("--evidence-dir", required=True, help="待分类证据目录")
    parser.add_argument("--output-dir", required=True, help="输出目录")
    parser.add_argument("--handbooks", nargs="*", help="审核手册 md 文件路径，不传则使用默认文件名")
    parser.add_argument("--threshold", type=float, default=DEFAULT_THRESHOLD, help="分类阈值，默认 3.5")
    return parser.parse_args()


def main():
    args = parse_args()
    script_dir = Path(__file__).resolve().parent
    evidence_dir = Path(args.evidence_dir).resolve()
    output_dir = Path(args.output_dir).resolve()

    if not evidence_dir.exists():
        print(f"[错误] 证据目录不存在：{evidence_dir}", file=sys.stderr)
        sys.exit(1)

    handbook_paths = [Path(p).resolve() for p in args.handbooks] if args.handbooks else default_handbook_paths(script_dir)
    rules = load_rules_from_handbooks(handbook_paths)
    if not rules:
        print("[错误] 未从手册中解析出任何规则。", file=sys.stderr)
        sys.exit(1)

    df_map = compute_keyword_df(rules)

    docs = [read_document(p) for p in iter_evidence_files(evidence_dir)]
    all_results: Dict[str, List[MatchResult]] = {}
    for doc in docs:
        all_results[str(doc.path)] = classify_document(doc, rules, df_map, args.threshold)

    write_results(output_dir, rules, docs, all_results)

    classified = sum(1 for v in all_results.values() if v)
    print(f"[完成] 共处理 {len(docs)} 个文件，命中分类 {classified} 个文件。")
    print(f"[输出目录] {output_dir}")
    print("已生成：classification_detail.csv / classification_summary.json / classification_by_control.csv / compiled_rules.csv / report.md")


if __name__ == "__main__":
    main()
